#!/usr/bin/env python3
"""Generate Unit Mix Summary table as a formatted Word document."""

import argparse
import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd elements
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_vertical_align(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_inches * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


def generate_table(data, output_path):
    doc = Document()

    # Narrow margins so table fits on one page
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    rows_data = data['rows']
    headers = [
        'UNIT TYPE',
        '%-INVENTORY',
        'MIN-SIZE (SF)',
        'MAX-SIZE (SF)',
        'AVERAGE-SIZE',
        'MEDIAN-SIZE',
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=6)
    table.style = 'Table Grid'

    # Column widths (inches): unit type wider, stats columns equal
    col_widths = [1.8, 1.1, 1.1, 1.1, 1.1, 1.1]

    # Header row
    header_row = table.rows[0]
    for i, (cell, header) in enumerate(zip(header_row.cells, headers)):
        set_cell_shading(cell, '1A1A1A')
        set_cell_vertical_align(cell, 'center')
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(8.5)
        run.font.name = 'Calibri'

    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        values = [
            row_data['unit_type'],
            row_data['pct_inventory'],
            str(row_data['min_size']),
            str(row_data['max_size']),
            str(row_data['avg_size']),
            str(row_data['median_size']),
        ]
        alt_bg = 'F2F2F2' if row_idx % 2 == 1 else 'FFFFFF'

        for col_idx, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_vertical_align(cell, 'center')
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(value)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

            if col_idx == 0:
                # Unit type column: yellow background, bold
                set_cell_shading(cell, 'FFFF00')
                run.font.bold = True
            else:
                set_cell_shading(cell, alt_bg)

    # Apply column widths after all rows created
    for col_idx, width in enumerate(col_widths):
        set_col_width(table, col_idx, width)

    doc.save(output_path)
    print(f"Saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Generate Unit Mix Summary Word doc')
    parser.add_argument('--data', help='JSON string with rows array')
    parser.add_argument('--output', default='Unit Mix Summary.docx', help='Output file path')
    args = parser.parse_args()

    if args.data:
        data = json.loads(args.data)
    else:
        data = json.loads(sys.stdin.read())

    generate_table(data, args.output)


if __name__ == '__main__':
    main()
